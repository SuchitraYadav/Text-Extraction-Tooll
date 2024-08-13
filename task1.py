import os
import pytesseract
from PIL import Image, ImageEnhance, ImageFilter, ImageOps
from pdf2image import convert_from_path
import docx
import re
import logging
from fpdf import FPDF
import io
import cv2
from spellchecker import SpellChecker


logging.basicConfig(
    filename="extraction.log",
    level=logging.DEBUG,  
    format='%(asctime)s - %(levelname)s - %(message)s',
    filemode='w'  
)


pytesseract.pytesseract.tesseract_cmd = r'C:\\Users\\syadav18\\AppData\\Local\\Programs\\Tesseract-OCR\\tesseract.exe'  # Update this path

# Preprocess the image
def preprocess_image(image):
    logging.debug("Starting image preprocessing")
    image = image.convert('L')  # Convert to grayscale
    enhancer = ImageEnhance.Contrast(image)
    image = enhancer.enhance(2)  # Increase contrast
    image = image.point(lambda x: 0 if x < 128 else 255, '1')  # Apply threshold
    logging.info("Preprocessed the image")
    return image

# Extract text from an image file
def extract_text_from_image(image_path):
    try:
        logging.info(f"Extracting text from image: {image_path}")
        image = Image.open(image_path)
        image = preprocess_image(image)
        text = pytesseract.image_to_string(image, lang='eng')
        logging.debug(f"Extracted text from image: {text[:100]}...")  
        return text
    except Exception as e:
        logging.error(f"Failed to extract text from image {image_path}: {e}", exc_info=True)
        return ""

# Extract text from a PDF file
def extract_text_from_pdf(pdf_path):
    try:
        logging.info(f"Extracting text from PDF: {pdf_path}")
        text = ""
        images = convert_from_path(pdf_path, poppler_path=r'C:\Users\syadav18\Desktop\Ml tasks\Task 1\Release-24.02.0-0\poppler-24.02.0\Library\bin')
        for i, image in enumerate(images):
            logging.debug(f"Processing page {i + 1} of {pdf_path}")
            image = preprocess_image(image)
            text += pytesseract.image_to_string(image, lang='eng')
        logging.debug(f"Extracted text from PDF: {text[:100]}...")  
        return text
    except Exception as e:
        logging.error(f"Failed to extract text from PDF {pdf_path}: {e}", exc_info=True)
        return ""

# Extract images from a DOCX file and return the text extracted from them
def extract_images_from_docx(doc):
    text = ""
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            try:
                logging.info("Extracting text from image in DOCX")
                image = Image.open(io.BytesIO(rel.target_part.blob))
                image = preprocess_image(image)
                text += pytesseract.image_to_string(image, lang='eng')
            except Exception as e:
                logging.error(f"Failed to extract text from an image in DOCX: {e}", exc_info=True)
    return text

# Extract text from a DOCX file
def extract_text_from_docx(docx_path):
    try:
        logging.info(f"Extracting text from DOCX: {docx_path}")
        doc = docx.Document(docx_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        text += extract_images_from_docx(doc)
        logging.debug(f"Extracted text from DOCX: {text[:100]}...")  
        return text
    except Exception as e:
        logging.error(f"Failed to extract text from DOCX {docx_path}: {e}", exc_info=True)
        return ""

# Extract text from a TXT file
def extract_text_from_txt(txt_path):
    try:
        logging.info(f"Extracting text from TXT: {txt_path}")
        with open(txt_path, 'r', encoding='utf-8') as file:
            text = file.read()
        logging.debug(f"Extracted text from TXT: {text[:100]}...")  
        return text
    except Exception as e:
        logging.error(f"Failed to extract text from TXT {txt_path}: {e}", exc_info=True)
        return ""

# Identify file type and extract text accordingly
def extract_text(file_path):
    try:
        logging.info(f"Identifying file type for extraction: {file_path}")
        if file_path.lower().endswith('.pdf'):
            return extract_text_from_pdf(file_path)
        elif file_path.lower().endswith('.docx'):
            return extract_text_from_docx(file_path)
        elif file_path.lower().endswith('.txt'):
            return extract_text_from_txt(file_path)
        elif file_path.lower().endswith(('.png', '.jpg', '.jpeg')):
            return extract_text_from_image(file_path)
        else:
            logging.error(f"Unsupported file type: {file_path}")
            return ""
    except Exception as e:
        logging.error(f"Failed to extract text from {file_path}: {e}", exc_info=True)
        return ""

# Validation method to check if the extracted words are correct
def validate_text(text):
    logging.info("Validating the extracted words")
    words = text.split()
    correct_words = []
    for word in words:
        if re.match("^[A-Za-z0-9_-]*$", word): 
            correct_words.append(word)
        else:
            logging.warning(f"Potentially incorrect word: {word}")
    validated_text = " ".join(correct_words)
    logging.debug(f"Validated text: {validated_text[:100]}...")  
    return validated_text

# Save the extracted text to a PDF file
def save_text_to_pdf(text, output_pdf_path):
    try:
        logging.info(f"Saving extracted text to PDF: {output_pdf_path}")
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)
        
        for line in text.split('\n'):
            pdf.multi_cell(0, 10, line)
        
        pdf.output(output_pdf_path)
        logging.info(f"Extracted text saved to {output_pdf_path}")
    except Exception as e:
        logging.error(f"Failed to save text to PDF {output_pdf_path}: {e}", exc_info=True)

# Save the extracted text to a DOCX file
def save_text_to_docx(text, output_docx_path):
    try:
        logging.info(f"Saving extracted text to DOCX: {output_docx_path}")
        doc = docx.Document()
        doc.add_paragraph(text)
        doc.save(output_docx_path)
        logging.info(f"Extracted text saved to {output_docx_path}")
    except Exception as e:
        logging.error(f"Failed to save text to DOCX {output_docx_path}: {e}", exc_info=True)

if __name__ == "__main__":
    try:
        logging.info("Script started")
        file_path = 'Input_file.pdf'  
        extracted_text = extract_text(file_path)
        validated_text = validate_text(extracted_text)
        
        output_pdf_path = 'pdf_text_output.pdf'  
        save_text_to_pdf(validated_text, output_pdf_path)
        
        output_docx_path = 'doc_text_output.docx'  
        save_text_to_docx(validated_text, output_docx_path)
        
        logging.info(f"Extracted text saved to {output_pdf_path} and {output_docx_path}")
    except Exception as e:
        logging.error(f"Unhandled exception in main: {e}", exc_info=True)
